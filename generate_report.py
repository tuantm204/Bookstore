#!/usr/bin/env python3
"""
Generate Technical Report for BookStore Microservices System
Output: BookStore_Technical_Report.docx (8-12 pages)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ============================================================
# STYLES SETUP
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# Heading styles
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        hs.font.size = Pt(18)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif i == 2:
        hs.font.size = Pt(15)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_paragraph(text, bold=False, italic=False, size=None, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2F5496")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D6E4F0")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    # Add shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
    p._p.get_or_add_pPr().append(shading)
    return p

def add_figure_caption(text, number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Hình {number}: {text}')
    run.bold = True
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(12)

def add_table_caption(text, number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Bảng {number}: {text}')
    run.bold = True
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)

def page_break():
    doc.add_page_break()

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_paragraph('BÁO CÁO KỸ THUẬT', bold=True, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_paragraph('HỆ THỐNG BOOKSTORE MICROSERVICES', bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_paragraph('Ứng dụng kiến trúc Microservices trong xây dựng', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, italic=False)
add_paragraph('hệ thống quản lý nhà sách trực tuyến', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(4):
    doc.add_paragraph()

add_paragraph('Công nghệ: Python / Django / Django REST Framework / MySQL / Docker', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_paragraph(f'Ngày: {datetime.date.today().strftime("%d/%m/%Y")}', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_paragraph('MỤC LỤC', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

toc_items = [
    ('1.', 'Giới thiệu', 3),
    ('1.1', 'Mục tiêu dự án', 3),
    ('1.2', 'Phạm vi hệ thống', 3),
    ('2.', 'Kiến trúc hệ thống', 4),
    ('2.1', 'Tổng quan kiến trúc Microservices', 4),
    ('2.2', 'Sơ đồ kiến trúc tổng thể', 4),
    ('2.3', 'Docker Network Topology', 5),
    ('3.', 'Công nghệ sử dụng', 5),
    ('4.', 'Chi tiết các Microservices', 6),
    ('4.1', 'Danh sách Services', 6),
    ('4.2', 'Giao tiếp giữa các Services', 6),
    ('5.', 'Thiết kế cơ sở dữ liệu', 7),
    ('5.1', 'Database per Service Pattern', 7),
    ('5.2', 'Data Models', 7),
    ('6.', 'API Gateway', 9),
    ('6.1', 'Chức năng', 9),
    ('6.2', 'Bảng định tuyến', 9),
    ('7.', 'API Documentation', 10),
    ('7.1', 'RESTful API Endpoints', 10),
    ('7.2', 'Ví dụ Request/Response', 10),
    ('8.', 'Luồng nghiệp vụ chính', 11),
    ('8.1', 'Luồng đặt hàng', 11),
    ('8.2', 'Luồng đăng ký khách hàng', 11),
    ('9.', 'Triển khai & Vận hành', 12),
    ('9.1', 'Docker Compose', 12),
    ('9.2', 'Hướng dẫn cài đặt', 12),
    ('10.', 'Đánh giá & Hướng phát triển', 13),
    ('10.1', 'Ưu điểm', 13),
    ('10.2', 'Hạn chế & Hướng cải thiện', 13),
    ('11.', 'Kết luận', 14),
]

for num, title, pg in toc_items:
    p = doc.add_paragraph()
    indent = '    ' if '.' in num and not num.endswith('.') else ''
    run = p.add_run(f'{indent}{num}  {title}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    if num.endswith('.'):
        run.bold = True
    # Add tab and page
    tab_run = p.add_run(f'\t{pg}')
    tab_run.font.name = 'Times New Roman'
    tab_run.font.size = Pt(13)
    p.paragraph_format.space_after = Pt(2)

page_break()

# ============================================================
# CHAPTER 1: INTRODUCTION
# ============================================================
doc.add_heading('1. Giới thiệu', level=1)

doc.add_heading('1.1 Mục tiêu dự án', level=2)
add_paragraph(
    'Dự án BookStore Microservices được xây dựng với mục tiêu phát triển một hệ thống '
    'quản lý nhà sách trực tuyến hoàn chỉnh, áp dụng kiến trúc Microservices hiện đại. '
    'Hệ thống cho phép quản lý toàn bộ quy trình kinh doanh của một nhà sách, từ quản lý '
    'sản phẩm, khách hàng, đặt hàng, thanh toán đến vận chuyển và đánh giá.'
)
add_paragraph('Các mục tiêu cụ thể bao gồm:')
add_bullet('Xây dựng hệ thống theo kiến trúc Microservices với các service độc lập')
add_bullet('Mỗi service có database riêng (Database per Service pattern)')
add_bullet('Triển khai bằng Docker Compose, đảm bảo tính nhất quán môi trường')
add_bullet('Cung cấp API Gateway làm điểm truy cập duy nhất cho client')
add_bullet('Hỗ trợ giao tiếp giữa các service qua REST API nội bộ')
add_bullet('Cung cấp giao diện Web UI để tương tác với hệ thống')

doc.add_heading('1.2 Phạm vi hệ thống', level=2)
add_paragraph(
    'Hệ thống bao gồm 11 microservice độc lập và 1 API Gateway, mỗi service đảm nhận '
    'một nghiệp vụ riêng biệt trong quy trình kinh doanh nhà sách:'
)

scope_data = [
    ['👤 Customer', 'Quản lý thông tin khách hàng, tự động tạo giỏ hàng'],
    ['👨‍💼 Staff', 'Quản lý nhân viên với phân quyền vai trò'],
    ['👔 Manager', 'Quản lý cấp quản lý theo phòng ban'],
    ['📂 Catalog', 'Phân loại sách theo danh mục/thể loại'],
    ['📖 Book', 'Quản lý sách, giá cả và tồn kho'],
    ['🛒 Cart', 'Giỏ hàng, quản lý sản phẩm đã chọn'],
    ['📦 Order', 'Xử lý đặt hàng, điều phối thanh toán & vận chuyển'],
    ['💳 Payment', 'Quản lý thanh toán (credit card, cash, transfer)'],
    ['🚚 Shipment', 'Quản lý vận chuyển, theo dõi đơn hàng'],
    ['💬 Comment', 'Bình luận và đánh giá sách (1-5 sao)'],
    ['🤖 Recommender', 'Gợi ý sách thông minh cho khách hàng'],
]
add_table_caption('Phạm vi chức năng các service', 1)
add_styled_table(['Service', 'Chức năng'], scope_data, [4, 12])

page_break()

# ============================================================
# CHAPTER 2: SYSTEM ARCHITECTURE
# ============================================================
doc.add_heading('2. Kiến trúc hệ thống', level=1)

doc.add_heading('2.1 Tổng quan kiến trúc Microservices', level=2)
add_paragraph(
    'Hệ thống được thiết kế theo kiến trúc Microservices, trong đó mỗi service là một '
    'ứng dụng Django độc lập với database riêng, chạy trong Docker container riêng biệt. '
    'Tất cả các service giao tiếp với nhau thông qua REST API qua mạng Docker nội bộ '
    '(bookstore-net).'
)
add_paragraph('Các nguyên tắc thiết kế chính được áp dụng:')

principles = [
    ['Single Responsibility', 'Mỗi service chỉ quản lý một domain nghiệp vụ duy nhất'],
    ['Database per Service', '11 database riêng biệt, không chia sẻ schema giữa các service'],
    ['API Gateway Pattern', 'Mọi request từ client đều đi qua một entry point duy nhất (port 8000)'],
    ['Service Orchestration', 'Order Service đóng vai trò orchestrator, điều phối nhiều service'],
    ['Loose Coupling', 'Các service liên kết qua ID (customer_id, book_id), không FK trực tiếp'],
    ['Independent Deployment', 'Mỗi service có Dockerfile riêng, build và deploy độc lập'],
]
add_table_caption('Nguyên tắc thiết kế Microservices', 2)
add_styled_table(['Nguyên tắc', 'Cách áp dụng'], principles, [5, 11])

doc.add_heading('2.2 Sơ đồ kiến trúc tổng thể', level=2)
add_paragraph(
    'Sơ đồ dưới đây mô tả kiến trúc tổng thể của hệ thống BookStore Microservices, '
    'bao gồm 3 tầng chính: Client Layer, Microservices Layer và Database Layer.'
)

# Architecture diagram as ASCII art in code block
arch_text = """
    ┌─────────────────────────────────────────────────────────┐
    │                   🌐 CLIENT LAYER                        │
    │         Web Browser (Bootstrap 5)  |  Postman            │
    └──────────────────────┬──────────────────────────────────┘
                           │ HTTP :8000
    ┌──────────────────────▼──────────────────────────────────┐
    │              ⚡ API GATEWAY (Django :8000)                │
    │         URL Proxy Routing + Web UI (HTML Templates)      │
    └──┬──┬──┬──┬──┬──┬──┬──┬──┬──┬──┬────────────────────────┘
       │  │  │  │  │  │  │  │  │  │  │
    ┌──▼──▼──▼──▼──▼──▼──▼──▼──▼──▼──▼────────────────────────┐
    │              🔧 MICROSERVICES LAYER                       │
    │                                                          │
    │  Customer  Staff  Manager  Catalog  Book                 │
    │   :8001    :8002   :8003    :8004   :8005                │
    │                                                          │
    │  Cart    Order   Payment  Shipment                       │
    │  :8006   :8007    :8008    :8009                         │
    │                                                          │
    │  Comment   Recommender                                   │
    │   :8010      :8011                                       │
    └──────────────────────┬──────────────────────────────────┘
                           │
    ┌──────────────────────▼──────────────────────────────────┐
    │              🗄️ DATABASE LAYER                            │
    │            MySQL 8.0 (11 databases)                      │
    │  customer_db | staff_db | manager_db | catalog_db        │
    │  book_db | cart_db | order_db | pay_db                   │
    │  ship_db | comment_db | recommend_db                     │
    └─────────────────────────────────────────────────────────┘"""

add_code_block(arch_text)
add_figure_caption('Sơ đồ kiến trúc tổng thể hệ thống BookStore Microservices', 1)

doc.add_heading('2.3 Docker Network Topology', level=2)
add_paragraph(
    'Toàn bộ 13 container (11 microservices + 1 API Gateway + 1 MySQL) đều nằm trên cùng '
    'một Docker network có tên "bookstore-net" (bridge driver). Các service giao tiếp với '
    'nhau bằng container name (DNS nội bộ Docker). MySQL sử dụng healthcheck để đảm bảo '
    'sẵn sàng trước khi các service khởi động.'
)

docker_text = """
    ┌────────────────── Docker Network: bookstore-net ──────────────────┐
    │                                                                   │
    │  api-gateway:8000 ──→ customer-service:8001                       │
    │                   ──→ staff-service:8002                          │
    │                   ──→ manager-service:8003                        │
    │                   ──→ catalog-service:8004                        │
    │                   ──→ book-service:8005                           │
    │                   ──→ cart-service:8006                           │
    │                   ──→ order-service:8007                          │
    │                   ──→ pay-service:8008                            │
    │                   ──→ ship-service:8009                           │
    │                   ──→ comment-rate-service:8010                   │
    │                   ──→ recommender-ai-service:8011                 │
    │                                                                   │
    │  All services ──→ mysql:3306 (healthcheck: mysqladmin ping)       │
    └───────────────────────────────────────────────────────────────────┘"""

add_code_block(docker_text)
add_figure_caption('Docker Network Topology', 2)

page_break()

# ============================================================
# CHAPTER 3: TECHNOLOGY STACK
# ============================================================
doc.add_heading('3. Công nghệ sử dụng', level=1)

add_paragraph(
    'Hệ thống BookStore Microservices sử dụng các công nghệ hiện đại, phổ biến '
    'trong phát triển web backend và triển khai containerized applications.'
)

tech_data = [
    ['Ngôn ngữ lập trình', 'Python', '3.11', 'Ngôn ngữ chính cho tất cả services'],
    ['Backend Framework', 'Django', '4.2.11', 'Web framework cho mỗi microservice'],
    ['REST API', 'Django REST Framework', '3.15.1', 'Xây dựng RESTful API'],
    ['Cơ sở dữ liệu', 'MySQL', '8.0', 'RDBMS cho 11 databases'],
    ['DB Client', 'mysqlclient', '2.2.4', 'Python MySQL connector'],
    ['HTTP Client', 'requests', '2.31.0', 'Giao tiếp REST giữa các service'],
    ['CORS', 'django-cors-headers', '4.3.1', 'Xử lý Cross-Origin requests'],
    ['Frontend', 'HTML5 + Bootstrap 5', '5.x', 'Giao diện Web UI responsive'],
    ['JavaScript', 'Fetch API (Vanilla)', 'ES6+', 'Gọi API từ frontend'],
    ['Container Runtime', 'Docker', '20.x+', 'Container hóa các service'],
    ['Orchestration', 'Docker Compose', 'v3.8', 'Quản lý multi-container'],
]
add_table_caption('Công nghệ sử dụng trong hệ thống', 3)
add_styled_table(['Thành phần', 'Công nghệ', 'Phiên bản', 'Mô tả'], tech_data, [3.5, 4, 2, 6.5])

add_paragraph(
    'Toàn bộ hệ thống được đóng gói trong Docker containers, sử dụng Docker Compose '
    'để orchestrate. Mỗi microservice có Dockerfile riêng dựa trên Python 3.11-slim image, '
    'cài đặt các system dependencies (libmysqlclient-dev, gcc) và Python packages qua pip.'
)

tech_text = """
    ┌─────────────────────────────────────────────────────────────────┐
    │                    🏗️ TECHNOLOGY STACK                           │
    │                                                                 │
    │  ┌────────────┐  ┌─────────────┐  ┌────────────────────────┐   │
    │  │  Frontend   │  │ API Gateway │  │    Backend (×11)       │   │
    │  ├────────────┤  ├─────────────┤  ├────────────────────────┤   │
    │  │ HTML5/CSS3 │  │ Django 4.2  │  │ Django 4.2             │   │
    │  │ Bootstrap 5│  │ SQLite      │  │ DRF 3.15.1             │   │
    │  │ JavaScript │  │ Requests    │  │ MySQL 8.0              │   │
    │  │ Fetch API  │  │ URL Proxy   │  │ mysqlclient            │   │
    │  └────────────┘  └─────────────┘  └────────────────────────┘   │
    │                                                                 │
    │  ┌────────────┐  ┌─────────────┐  ┌────────────────────────┐   │
    │  │ Infra      │  │ Language    │  │ Patterns               │   │
    │  ├────────────┤  ├─────────────┤  ├────────────────────────┤   │
    │  │ Docker     │  │ Python 3.11 │  │ API Gateway Pattern    │   │
    │  │ Compose    │  │             │  │ Database per Service   │   │
    │  │ MySQL 8.0  │  │             │  │ Service Orchestration  │   │
    │  └────────────┘  └─────────────┘  └────────────────────────┘   │
    └─────────────────────────────────────────────────────────────────┘"""

add_code_block(tech_text)
add_figure_caption('Technology Stack Overview', 3)

page_break()

# ============================================================
# CHAPTER 4: MICROSERVICES DETAIL
# ============================================================
doc.add_heading('4. Chi tiết các Microservices', level=1)

doc.add_heading('4.1 Danh sách Services', level=2)
add_paragraph(
    'Hệ thống gồm 12 thành phần (11 microservices + 1 API Gateway), mỗi thành phần '
    'chạy trong một Docker container riêng biệt:'
)

services_data = [
    ['1', 'Customer Service', '8001', 'customer_db', 'Quản lý khách hàng, auto-tạo giỏ hàng'],
    ['2', 'Staff Service', '8002', 'staff_db', 'Quản lý nhân viên nhà sách'],
    ['3', 'Manager Service', '8003', 'manager_db', 'Quản lý quản lý theo phòng ban'],
    ['4', 'Catalog Service', '8004', 'catalog_db', 'Quản lý danh mục/thể loại sách'],
    ['5', 'Book Service', '8005', 'book_db', 'Quản lý sách, giá, tồn kho'],
    ['6', 'Cart Service', '8006', 'cart_db', 'Giỏ hàng + Cart Items'],
    ['7', 'Order Service', '8007', 'order_db', 'Đặt hàng, orchestrate payment & shipment'],
    ['8', 'Payment Service', '8008', 'pay_db', 'Thanh toán đơn hàng'],
    ['9', 'Shipment Service', '8009', 'ship_db', 'Vận chuyển, tracking đơn hàng'],
    ['10', 'Comment & Rate', '8010', 'comment_db', 'Bình luận & đánh giá sách'],
    ['11', 'Recommender AI', '8011', 'recommend_db', 'Gợi ý sách thông minh'],
    ['12', 'API Gateway', '8000', 'SQLite', 'Proxy routing + Web UI'],
]
add_table_caption('Danh sách các Microservices', 4)
add_styled_table(['#', 'Service', 'Port', 'Database', 'Mô tả'], services_data, [1, 3.5, 1.5, 2.5, 7.5])

doc.add_heading('4.2 Giao tiếp giữa các Services (Inter-Service Communication)', level=2)
add_paragraph(
    'Các microservice giao tiếp với nhau qua REST API nội bộ (synchronous HTTP) '
    'thông qua Docker network. Dưới đây là danh sách các kết nối inter-service:'
)

comm_data = [
    ['①', 'Customer → Cart', 'POST /carts/', 'Tự động tạo giỏ hàng rỗng khi đăng ký customer mới'],
    ['②', 'Cart → Book', 'GET /books/{id}/', 'Kiểm tra sách tồn tại khi thêm vào giỏ'],
    ['③', 'Order → Cart', 'GET /carts/customer/{id}/', 'Lấy tất cả items trong giỏ hàng'],
    ['④', 'Order → Book', 'GET /books/{id}/', 'Lấy giá sách để tính tổng đơn hàng'],
    ['⑤', 'Order → Payment', 'POST /payments/', 'Tự động tạo bản ghi thanh toán (pending)'],
    ['⑥', 'Order → Shipment', 'POST /shipments/', 'Tự động tạo bản ghi vận chuyển'],
    ['⑦', 'Recommender → Book', 'GET /books/', 'Lấy danh sách sách để gợi ý ngẫu nhiên'],
]
add_table_caption('Inter-Service Communication', 5)
add_styled_table(['#', 'Kết nối', 'API Call', 'Mô tả'], comm_data, [1, 3.5, 4.5, 7])

comm_text = """
    Customer ──POST /carts/──────────────────────►  Cart
    Cart     ──GET /books/{id}/──────────────────►  Book
    Order    ──GET /carts/customer/{id}/─────────►  Cart
    Order    ──GET /books/{id}/──────────────────►  Book
    Order    ──POST /payments/───────────────────►  Payment
    Order    ──POST /shipments/──────────────────►  Shipment
    Recommender ──GET /books/────────────────────►  Book"""

add_code_block(comm_text)
add_figure_caption('Sơ đồ giao tiếp giữa các Services', 4)

page_break()

# ============================================================
# CHAPTER 5: DATABASE DESIGN
# ============================================================
doc.add_heading('5. Thiết kế cơ sở dữ liệu', level=1)

doc.add_heading('5.1 Database per Service Pattern', level=2)
add_paragraph(
    'Hệ thống áp dụng mô hình "Database per Service" — mỗi microservice sở hữu một '
    'database riêng biệt trong cùng MySQL server. Không có service nào truy cập trực tiếp '
    'vào database của service khác. Việc liên kết dữ liệu giữa các service được thực hiện '
    'thông qua ID tham chiếu (customer_id, book_id, order_id) và REST API calls.'
)

db_data = [
    ['customer_db', 'Customer Service', 'customers_customer'],
    ['staff_db', 'Staff Service', 'staffs_staff'],
    ['manager_db', 'Manager Service', 'managers_manager'],
    ['catalog_db', 'Catalog Service', 'catalogs_catalog'],
    ['book_db', 'Book Service', 'books_book'],
    ['cart_db', 'Cart Service', 'carts_cart, carts_cartitem'],
    ['order_db', 'Order Service', 'orders_order, orders_orderitem'],
    ['pay_db', 'Payment Service', 'payments_payment'],
    ['ship_db', 'Shipment Service', 'shipments_shipment'],
    ['comment_db', 'Comment Service', 'comments_comment'],
    ['recommend_db', 'Recommender Service', 'recommendations_recommendation'],
]
add_table_caption('Database mapping cho các services', 6)
add_styled_table(['Database', 'Service', 'Tables'], db_data, [3.5, 4, 8.5])

db_text = """
    MySQL 8.0 Server (port 3306 internal / 3307 external)
    ├── customer_db    ← Customer Service :8001
    ├── staff_db       ← Staff Service :8002
    ├── manager_db     ← Manager Service :8003
    ├── catalog_db     ← Catalog Service :8004
    ├── book_db        ← Book Service :8005
    ├── cart_db        ← Cart Service :8006
    ├── order_db       ← Order Service :8007
    ├── pay_db         ← Payment Service :8008
    ├── ship_db        ← Shipment Service :8009
    ├── comment_db     ← Comment Service :8010
    └── recommend_db   ← Recommender Service :8011"""

add_code_block(db_text)
add_figure_caption('Database per Service Architecture', 5)

doc.add_heading('5.2 Data Models', level=2)
add_paragraph('Dưới đây là chi tiết các data model của từng service:')

# Customer model
add_paragraph('Customer Model:', bold=True, size=13)
add_styled_table(
    ['Field', 'Type', 'Ghi chú'],
    [
        ['id', 'AutoField', 'Primary key'],
        ['name', 'CharField(100)', 'Tên khách hàng'],
        ['email', 'EmailField', 'Unique - email khách hàng'],
        ['phone', 'CharField(20)', 'Số điện thoại'],
        ['address', 'TextField', 'Địa chỉ'],
        ['created_at', 'DateTimeField', 'Auto - thời gian tạo'],
    ], [3, 4, 9]
)

# Book model
add_paragraph('Book Model:', bold=True, size=13)
add_styled_table(
    ['Field', 'Type', 'Ghi chú'],
    [
        ['id', 'AutoField', 'Primary key'],
        ['title', 'CharField(200)', 'Tên sách'],
        ['author', 'CharField(100)', 'Tác giả'],
        ['price', 'DecimalField(10,2)', 'Giá sách'],
        ['stock', 'IntegerField', 'Số lượng tồn kho (default: 0)'],
        ['created_at', 'DateTimeField', 'Auto - thời gian tạo'],
    ], [3, 4, 9]
)

# Cart + CartItem
add_paragraph('Cart & CartItem Models:', bold=True, size=13)
add_styled_table(
    ['Model', 'Field', 'Type', 'Ghi chú'],
    [
        ['Cart', 'id', 'AutoField', 'Primary key'],
        ['Cart', 'customer_id', 'IntegerField', 'ID khách hàng (cross-service ref)'],
        ['Cart', 'created_at', 'DateTimeField', 'Auto'],
        ['CartItem', 'id', 'AutoField', 'Primary key'],
        ['CartItem', 'cart', 'ForeignKey(Cart)', 'FK → Cart'],
        ['CartItem', 'book_id', 'IntegerField', 'ID sách (cross-service ref)'],
        ['CartItem', 'quantity', 'IntegerField', 'Số lượng (default: 1)'],
    ], [2.5, 3, 4, 6.5]
)

# Order + OrderItem
add_paragraph('Order & OrderItem Models:', bold=True, size=13)
add_styled_table(
    ['Model', 'Field', 'Type', 'Ghi chú'],
    [
        ['Order', 'id', 'AutoField', 'Primary key'],
        ['Order', 'customer_id', 'IntegerField', 'ID khách hàng'],
        ['Order', 'total_amount', 'DecimalField(10,2)', 'Tổng tiền đơn hàng'],
        ['Order', 'status', 'CharField', 'pending/paid/shipped/delivered/cancelled'],
        ['Order', 'created_at', 'DateTimeField', 'Auto'],
        ['OrderItem', 'id', 'AutoField', 'Primary key'],
        ['OrderItem', 'order', 'ForeignKey(Order)', 'FK → Order'],
        ['OrderItem', 'book_id', 'IntegerField', 'ID sách'],
        ['OrderItem', 'quantity', 'IntegerField', 'Số lượng'],
        ['OrderItem', 'price', 'DecimalField(10,2)', 'Giá tại thời điểm mua'],
    ], [2.5, 3, 4, 6.5]
)

# Payment
add_paragraph('Payment Model:', bold=True, size=13)
add_styled_table(
    ['Field', 'Type', 'Ghi chú'],
    [
        ['id', 'AutoField', 'Primary key'],
        ['order_id', 'IntegerField', 'ID đơn hàng (cross-service ref)'],
        ['amount', 'DecimalField(10,2)', 'Số tiền thanh toán'],
        ['method', 'CharField', 'credit_card / cash / transfer'],
        ['status', 'CharField', 'pending / completed / failed'],
        ['created_at', 'DateTimeField', 'Auto'],
    ], [3, 4, 9]
)

# Shipment
add_paragraph('Shipment Model:', bold=True, size=13)
add_styled_table(
    ['Field', 'Type', 'Ghi chú'],
    [
        ['id', 'AutoField', 'Primary key'],
        ['order_id', 'IntegerField', 'ID đơn hàng'],
        ['address', 'TextField', 'Địa chỉ giao hàng'],
        ['status', 'CharField', 'processing / shipped / delivered'],
        ['tracking_number', 'CharField(100)', 'Mã vận đơn'],
        ['created_at', 'DateTimeField', 'Auto'],
    ], [3, 4, 9]
)

# Comment
add_paragraph('Comment Model:', bold=True, size=13)
add_styled_table(
    ['Field', 'Type', 'Ghi chú'],
    [
        ['id', 'AutoField', 'Primary key'],
        ['customer_id', 'IntegerField', 'ID khách hàng'],
        ['book_id', 'IntegerField', 'ID sách'],
        ['content', 'TextField', 'Nội dung bình luận'],
        ['rating', 'IntegerField', 'Đánh giá 1-5 sao (default: 5)'],
        ['created_at', 'DateTimeField', 'Auto'],
    ], [3, 4, 9]
)

# Recommendation
add_paragraph('Recommendation Model:', bold=True, size=13)
add_styled_table(
    ['Field', 'Type', 'Ghi chú'],
    [
        ['id', 'AutoField', 'Primary key'],
        ['customer_id', 'IntegerField', 'ID khách hàng'],
        ['book_id', 'IntegerField', 'ID sách được gợi ý'],
        ['score', 'FloatField', 'Điểm gợi ý (0.0 - 1.0)'],
        ['created_at', 'DateTimeField', 'Auto'],
    ], [3, 4, 9]
)

page_break()

# ============================================================
# CHAPTER 6: API GATEWAY
# ============================================================
doc.add_heading('6. API Gateway', level=1)

doc.add_heading('6.1 Chức năng', level=2)
add_paragraph(
    'API Gateway là thành phần trung tâm của hệ thống, đóng vai trò là điểm truy cập '
    'duy nhất (single entry point) cho tất cả client requests. Gateway được xây dựng bằng '
    'Django và thực hiện các chức năng sau:'
)
add_bullet('URL Proxy Routing: Chuyển tiếp request từ client đến đúng microservice tương ứng')
add_bullet('Web UI Hosting: Phục vụ các trang HTML (Bootstrap 5) cho quản lý hệ thống')
add_bullet('Service Abstraction: Che giấu topology nội bộ của các microservice khỏi client')
add_bullet('Unified Entry Point: Tất cả API đều truy cập qua cổng 8000')

add_paragraph(
    'Gateway sử dụng thư viện "requests" của Python để proxy các HTTP request. '
    'Khi nhận request tại /api/{service}/{path}, gateway sẽ tra cứu SERVICE_MAP '
    'để tìm URL nội bộ tương ứng và chuyển tiếp request.'
)

doc.add_heading('6.2 Bảng định tuyến (Routing Table)', level=2)

routing_data = [
    ['/api/customers/*', 'http://customer-service:8001/customers/*'],
    ['/api/staffs/*', 'http://staff-service:8002/staffs/*'],
    ['/api/managers/*', 'http://manager-service:8003/managers/*'],
    ['/api/catalogs/*', 'http://catalog-service:8004/catalogs/*'],
    ['/api/books/*', 'http://book-service:8005/books/*'],
    ['/api/carts/*', 'http://cart-service:8006/carts/*'],
    ['/api/cart-items/*', 'http://cart-service:8006/cart-items/*'],
    ['/api/orders/*', 'http://order-service:8007/orders/*'],
    ['/api/order-items/*', 'http://order-service:8007/order-items/*'],
    ['/api/payments/*', 'http://pay-service:8008/payments/*'],
    ['/api/shipments/*', 'http://ship-service:8009/shipments/*'],
    ['/api/comments/*', 'http://comment-service:8010/comments/*'],
    ['/api/recommendations/*', 'http://recommender-service:8011/recommendations/*'],
    ['/api/recommend/*', 'http://recommender-service:8011/recommend/*'],
]
add_table_caption('API Gateway Routing Table', 7)
add_styled_table(['Client Request', 'Internal Proxy Target'], routing_data, [6, 10])

page_break()

# ============================================================
# CHAPTER 7: API DOCUMENTATION
# ============================================================
doc.add_heading('7. API Documentation', level=1)

doc.add_heading('7.1 RESTful API Endpoints', level=2)
add_paragraph(
    'Tất cả API endpoints tuân theo chuẩn RESTful. Base URL: http://localhost:8000/api. '
    'Mỗi service cung cấp các operations CRUD cơ bản (GET, POST, PUT, PATCH, DELETE).'
)

api_data = [
    ['Customer', 'GET/POST /api/customers/', '/api/customers/{id}/'],
    ['Staff', 'GET/POST /api/staffs/', '/api/staffs/{id}/'],
    ['Manager', 'GET/POST /api/managers/', '/api/managers/{id}/'],
    ['Catalog', 'GET/POST /api/catalogs/', '/api/catalogs/{id}/'],
    ['Book', 'GET/POST /api/books/', '/api/books/{id}/'],
    ['Cart', 'GET /api/carts/', '/api/carts/{id}/'],
    ['CartItem', 'GET/POST /api/cart-items/', '/api/cart-items/{id}/'],
    ['Order', 'GET/POST /api/orders/', '/api/orders/{id}/'],
    ['OrderItem', 'GET /api/order-items/', '/api/order-items/{id}/'],
    ['Payment', 'GET/POST /api/payments/', '/api/payments/{id}/'],
    ['Shipment', 'GET/POST /api/shipments/', '/api/shipments/{id}/'],
    ['Comment', 'GET/POST /api/comments/', '/api/comments/{id}/'],
    ['Recommendation', 'GET/POST /api/recommendations/', '/api/recommendations/{id}/'],
]
add_table_caption('RESTful API Endpoints', 8)
add_styled_table(['Service', 'List & Create', 'Detail (GET/PUT/PATCH/DELETE)'], api_data, [3, 5.5, 7.5])

add_paragraph('Custom Endpoints (ngoài CRUD cơ bản):', bold=True)
custom_data = [
    ['GET /api/carts/customer/{customer_id}/', 'Lấy giỏ hàng theo Customer ID'],
    ['GET /api/recommend/', 'Gợi ý sách ngẫu nhiên (AI Recommendation)'],
]
add_styled_table(['Endpoint', 'Mô tả'], custom_data, [7, 9])

doc.add_heading('7.2 Ví dụ Request/Response', level=2)
add_paragraph('Tạo khách hàng mới:', bold=True)
add_code_block(
    'POST /api/customers/\n'
    'Content-Type: application/json\n\n'
    '{\n'
    '    "name": "Nguyen Van A",\n'
    '    "email": "a@gmail.com",\n'
    '    "phone": "0901234567",\n'
    '    "address": "Ha Noi"\n'
    '}\n\n'
    'Response: 201 Created\n'
    '{\n'
    '    "id": 1,\n'
    '    "name": "Nguyen Van A",\n'
    '    "email": "a@gmail.com",\n'
    '    "phone": "0901234567",\n'
    '    "address": "Ha Noi",\n'
    '    "created_at": "2026-03-09T10:00:00Z"\n'
    '}'
)

add_paragraph('Tạo đơn hàng (auto-orchestrate):', bold=True)
add_code_block(
    'POST /api/orders/\n'
    'Content-Type: application/json\n\n'
    '{\n'
    '    "customer_id": 1,\n'
    '    "address": "123 Nguyen Hue, HCM"\n'
    '}\n\n'
    'Response: 201 Created\n'
    '{\n'
    '    "id": 1,\n'
    '    "customer_id": 1,\n'
    '    "total_amount": "44.48",\n'
    '    "status": "pending",\n'
    '    "items": [\n'
    '        {"book_id": 1, "quantity": 2, "price": "15.99"},\n'
    '        {"book_id": 2, "quantity": 1, "price": "12.50"}\n'
    '    ]\n'
    '}\n'
    '→ Auto-created: Payment (pending) + Shipment (processing)'
)

page_break()

# ============================================================
# CHAPTER 8: BUSINESS FLOW
# ============================================================
doc.add_heading('8. Luồng nghiệp vụ chính', level=1)

doc.add_heading('8.1 Luồng đặt hàng (Order Creation Flow)', level=2)
add_paragraph(
    'Đây là luồng nghiệp vụ phức tạp nhất trong hệ thống, thể hiện rõ mô hình '
    'Service Orchestration. Khi client gửi request tạo đơn hàng, Order Service đóng '
    'vai trò orchestrator và điều phối 4 service khác:'
)

order_flow = """
    Client ──POST /api/orders/──► Gateway ──► Order Service
                                                    │
                        ┌───────────────────────────┘
                        │
                Step 1: │──GET /carts/customer/{id}/──► Cart Service
                        │   ◄── {items: [{book_id, qty}]}
                        │
                Step 2: │──GET /books/{book_id}/──► Book Service (loop)
                        │   ◄── {price, stock}
                        │
                Step 3: │  Calculate: total = Σ(price × qty)
                        │  Save Order + OrderItems to order_db
                        │
                Step 4: │──POST /payments/──► Payment Service
                        │   ◄── {id, status: "pending"}
                        │
                Step 5: │──POST /shipments/──► Shipment Service
                        │   ◄── {id, status: "processing"}
                        │
                        └──► Response: 201 {order + items}"""

add_code_block(order_flow)
add_figure_caption('Sequence Diagram — Luồng đặt hàng', 6)

add_paragraph('Chi tiết các bước:')
add_bullet('Bước 1: Order Service gọi Cart Service để lấy danh sách items trong giỏ hàng của customer')
add_bullet('Bước 2: Với mỗi item, Order Service gọi Book Service để lấy giá sách hiện tại')
add_bullet('Bước 3: Tính tổng tiền (total_amount = Σ price × quantity), lưu Order và OrderItems')
add_bullet('Bước 4: Tự động tạo bản ghi Payment (status: pending, method: credit_card)')
add_bullet('Bước 5: Tự động tạo bản ghi Shipment (status: processing, address từ request)')

doc.add_heading('8.2 Luồng đăng ký khách hàng (Customer Registration Flow)', level=2)
add_paragraph(
    'Khi một khách hàng mới được tạo, Customer Service tự động gọi Cart Service '
    'để tạo một giỏ hàng rỗng cho khách hàng đó:'
)

cust_flow = """
    Client ──POST /api/customers/──► Gateway ──► Customer Service
                                                        │
                                    Save customer ◄─────┘
                                    to customer_db
                                          │
                                          │──POST /carts/──► Cart Service
                                          │   {customer_id: new_id}
                                          │   ◄── {id, customer_id}
                                          │
                                          └──► Response: 201 {customer data}"""

add_code_block(cust_flow)
add_figure_caption('Sequence Diagram — Luồng đăng ký khách hàng', 7)

add_paragraph('Tóm tắt toàn bộ Business Flow:', bold=True)
add_bullet('1. Tạo Customer → Auto-tạo Cart rỗng')
add_bullet('2. Tạo Books → Thêm sách vào hệ thống')
add_bullet('3. Add CartItem → Thêm sách vào giỏ (verify book tồn tại)')
add_bullet('4. Create Order → Fetch cart → Fetch prices → Create order → Auto Payment → Auto Shipment')
add_bullet('5. Add Comment → Đánh giá sách (1-5 sao)')
add_bullet('6. Get Recommendations → Gợi ý sách cho khách hàng')

page_break()

# ============================================================
# CHAPTER 9: DEPLOYMENT
# ============================================================
doc.add_heading('9. Triển khai & Vận hành', level=1)

doc.add_heading('9.1 Docker Compose', level=2)
add_paragraph(
    'Toàn bộ hệ thống được triển khai bằng Docker Compose (version 3.8). '
    'File docker-compose.yml định nghĩa 13 services, 1 network và 1 volume:'
)

compose_data = [
    ['mysql', 'mysql:8.0', '3307:3306', 'Database server với healthcheck'],
    ['customer-service', 'Build ./customer-service', '8001:8001', 'depends_on: mysql (healthy)'],
    ['staff-service', 'Build ./staff-service', '8002:8002', 'depends_on: mysql (healthy)'],
    ['manager-service', 'Build ./manager-service', '8003:8003', 'depends_on: mysql (healthy)'],
    ['catalog-service', 'Build ./catalog-service', '8004:8004', 'depends_on: mysql (healthy)'],
    ['book-service', 'Build ./book-service', '8005:8005', 'depends_on: mysql (healthy)'],
    ['cart-service', 'Build ./cart-service', '8006:8006', 'depends_on: mysql (healthy)'],
    ['order-service', 'Build ./order-service', '8007:8007', 'depends_on: mysql (healthy)'],
    ['pay-service', 'Build ./pay-service', '8008:8008', 'depends_on: mysql (healthy)'],
    ['ship-service', 'Build ./ship-service', '8009:8009', 'depends_on: mysql (healthy)'],
    ['comment-rate-service', 'Build ./comment-rate-service', '8010:8010', 'depends_on: mysql (healthy)'],
    ['recommender-ai-service', 'Build ./recommender-ai-service', '8011:8011', 'depends_on: mysql (healthy)'],
    ['api-gateway', 'Build ./api-gateway', '8000:8000', 'depends_on: all 11 services'],
]
add_table_caption('Docker Compose Services', 9)
add_styled_table(['Container', 'Image / Build', 'Port Mapping', 'Dependency'], compose_data, [3.5, 4, 2.5, 6])

add_paragraph(
    'MySQL sử dụng healthcheck (mysqladmin ping) với interval 10s, timeout 5s, retries 10, '
    'start_period 30s. Tất cả microservices chỉ khởi động khi MySQL healthy. '
    'Các service có restart: on-failure để tự khởi động lại nếu gặp lỗi.'
)

doc.add_heading('9.2 Hướng dẫn cài đặt', level=2)
add_paragraph('Yêu cầu hệ thống:', bold=True)
add_bullet('Docker ≥ 20.x và Docker Compose ≥ 2.x')
add_bullet('RAM ≥ 4GB (13 containers + MySQL)')
add_bullet('Cổng 8000-8011 và 3307 khả dụng')

add_paragraph('Các bước cài đặt:', bold=True)
add_code_block(
    '# Bước 1: Di chuyển vào thư mục\n'
    'cd bookstore-micro05\n\n'
    '# Bước 2: Build và khởi chạy\n'
    'docker-compose up --build -d\n\n'
    '# Bước 3: Kiểm tra trạng thái\n'
    'docker-compose ps\n\n'
    '# Bước 4: Chạy migrations (lần đầu)\n'
    'for svc in customer-service staff-service manager-service \\\n'
    '  catalog-service book-service cart-service order-service \\\n'
    '  pay-service ship-service comment-rate-service \\\n'
    '  recommender-ai-service; do\n'
    '  docker-compose exec $svc python manage.py makemigrations\n'
    '  docker-compose exec $svc python manage.py migrate\n'
    'done\n\n'
    '# Bước 5: Truy cập\n'
    '# Web UI:  http://localhost:8000\n'
    '# API:     http://localhost:8000/api/books/\n\n'
    '# Dừng hệ thống\n'
    'docker-compose down        # Giữ data\n'
    'docker-compose down -v     # Xóa data'
)

page_break()

# ============================================================
# CHAPTER 10: EVALUATION
# ============================================================
doc.add_heading('10. Đánh giá & Hướng phát triển', level=1)

doc.add_heading('10.1 Ưu điểm', level=2)

advantages = [
    ['Kiến trúc Microservices', 'Mỗi service độc lập, có thể phát triển, test và deploy riêng biệt. '
     'Team khác nhau có thể làm việc song song trên các service khác nhau.'],
    ['Database per Service', 'Dữ liệu được tách biệt hoàn toàn giữa các service, đảm bảo tính '
     'encapsulation và autonomy. Mỗi service có thể chọn database phù hợp.'],
    ['API Gateway', 'Client chỉ cần biết một endpoint duy nhất (port 8000), '
     'không cần biết topology nội bộ. Dễ dàng thêm authentication, rate limiting.'],
    ['Docker Compose', 'Toàn bộ hệ thống khởi động bằng 1 lệnh. Đảm bảo consistency '
     'giữa development, testing và production environments.'],
    ['Service Orchestration', 'Order Service thể hiện rõ pattern orchestration, '
     'điều phối nhiều service để hoàn thành nghiệp vụ phức tạp.'],
    ['RESTful API', 'Tất cả service tuân theo chuẩn REST, dễ tích hợp với bất kỳ client nào '
     '(web, mobile, third-party).'],
    ['Web UI', 'Cung cấp giao diện Bootstrap 5 responsive, cho phép thao tác CRUD trực quan '
     'mà không cần công cụ bên ngoài.'],
]
add_table_caption('Ưu điểm của hệ thống', 10)
add_styled_table(['Ưu điểm', 'Mô tả'], advantages, [4, 12])

doc.add_heading('10.2 Hạn chế & Hướng cải thiện', level=2)

limitations = [
    ['Synchronous Communication', 'Tất cả inter-service calls là synchronous REST, '
     'có thể gây cascading failures', 'Thêm Message Queue (RabbitMQ/Kafka) cho async'],
    ['No Authentication', 'API không có authentication/authorization', 
     'Thêm JWT/OAuth2 tại API Gateway'],
    ['No Service Discovery', 'Service URLs hardcoded trong code', 
     'Thêm Consul/Eureka cho dynamic discovery'],
    ['No Circuit Breaker', 'Không có cơ chế fallback khi service down', 
     'Thêm Circuit Breaker pattern (retry + fallback)'],
    ['No Centralized Logging', 'Logs phân tán trong các container', 
     'Thêm ELK Stack hoặc Loki cho centralized logging'],
    ['No Monitoring', 'Không có metrics/alerting', 
     'Thêm Prometheus + Grafana cho monitoring'],
    ['Simple Recommender', 'Gợi ý sách chỉ là random', 
     'Thêm ML-based recommendation engine'],
    ['No Data Consistency', 'Không có saga pattern cho distributed transactions', 
     'Implement Saga pattern hoặc Event Sourcing'],
]
add_table_caption('Hạn chế và hướng cải thiện', 11)
add_styled_table(['Hạn chế', 'Chi tiết', 'Hướng cải thiện'], limitations, [3.5, 5.5, 7])

page_break()

# ============================================================
# CHAPTER 11: CONCLUSION
# ============================================================
doc.add_heading('11. Kết luận', level=1)

add_paragraph(
    'Hệ thống BookStore Microservices đã được xây dựng thành công với kiến trúc microservices '
    'hoàn chỉnh, bao gồm 11 dịch vụ độc lập và 1 API Gateway. Hệ thống thể hiện rõ các '
    'nguyên tắc thiết kế microservices quan trọng:'
)

add_bullet('Single Responsibility: Mỗi service đảm nhận đúng một nghiệp vụ')
add_bullet('Database per Service: 11 database riêng biệt, đảm bảo data isolation')
add_bullet('API Gateway Pattern: Một entry point duy nhất cho client')
add_bullet('Service Orchestration: Order Service điều phối quy trình đặt hàng phức tạp')
add_bullet('Containerization: Toàn bộ hệ thống đóng gói trong Docker containers')

add_paragraph(
    'Hệ thống đã triển khai thành công trên Docker Compose với 13 containers hoạt động '
    'ổn định, bao gồm MySQL 8.0 với healthcheck, 11 Django microservices và 1 API Gateway '
    'với Web UI. Các tính năng chính hoạt động đầy đủ: quản lý khách hàng, sách, giỏ hàng, '
    'đặt hàng tự động (orchestration), thanh toán, vận chuyển, đánh giá và gợi ý sách.'
)

add_paragraph(
    'Dự án cũng đi kèm với bộ tài liệu đầy đủ bao gồm: API Documentation chi tiết cho '
    'tất cả endpoints, Architecture Diagrams (Mermaid) cho sơ đồ kiến trúc, và Postman '
    'Collection với 50+ requests sẵn sàng test. Đây là nền tảng tốt để tiếp tục phát triển '
    'thêm các tính năng nâng cao như authentication, message queue, monitoring và '
    'machine learning-based recommendations trong tương lai.'
)

doc.add_paragraph()
add_paragraph('— Hết —', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)

# ============================================================
# APPENDIX: File listing
# ============================================================
page_break()
doc.add_heading('Phụ lục: Cấu trúc thư mục dự án', level=1)

dir_text = """bookstore-micro05/
├── docker-compose.yml
├── init.sql
├── api-gateway/
│   ├── Dockerfile
│   ├── manage.py
│   ├── requirements.txt
│   ├── gateway_project/
│   │   ├── settings.py
│   │   └── urls.py
│   ├── gateway/
│   │   ├── views.py
│   │   └── urls.py
│   └── templates/
│       ├── index.html
│       ├── books.html
│       ├── cart.html
│       ├── customers.html
│       ├── orders.html
│       ├── payments.html
│       ├── shipments.html
│       ├── staffs.html
│       ├── managers.html
│       ├── catalogs.html
│       ├── comments.html
│       └── recommender.html
├── customer-service/
│   ├── Dockerfile
│   ├── manage.py
│   ├── requirements.txt
│   ├── customer_project/settings.py
│   └── customers/
│       ├── models.py
│       ├── serializers.py
│       ├── views.py
│       ├── urls.py
│       └── migrations/
├── staff-service/          (same structure)
├── manager-service/        (same structure)
├── catalog-service/        (same structure)
├── book-service/           (same structure)
├── cart-service/           (same structure)
├── order-service/          (same structure)
├── pay-service/            (same structure)
├── ship-service/           (same structure)
├── comment-rate-service/   (same structure)
├── recommender-ai-service/ (same structure)
├── API_DOCUMENTATION.md
├── ARCHITECTURE_DIAGRAM.md
├── BookStore_API.postman_collection.json
└── README.md"""

add_code_block(dir_text)

# ============================================================
# SAVE
# ============================================================
output_path = '/Users/truongmanhtuan/django_project/assgn05v1/bookstore-micro05/BookStore_Technical_Report.docx'
doc.save(output_path)
print(f'✅ Report saved: {output_path}')
print(f'📄 File generated successfully!')
